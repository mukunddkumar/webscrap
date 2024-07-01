import requests
from bs4 import BeautifulSoup
from docx import Document

# URL of the website
url = "https://insights.blackcoffer.com/rising-it-cities-and-their-impact-on-the-economy-environment-infrastructure-and-city-life-in-future/"

# Fetch the content from the URL
response = requests.get(url)
soup = BeautifulSoup(response.content, "html.parser")

# Extract the article content
article = soup.find("div", class_="td-post-content")
paragraphs = article.find_all("p")

# Create a new Document
doc = Document()
doc.add_heading("Rising IT Cities and Their Impact on the Economy, Environment, Infrastructure, and City Life in Future", 0)

# Add paragraphs to the Document
for paragraph in paragraphs:
    doc.add_paragraph(paragraph.get_text())

# Save the Document
file_path = "Rising_IT_Cities_Impact.docx"
doc.save(file_path)

print(f"The article has been saved to {file_path}")